from __future__ import annotations

import csv
import io
import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


class DocumentParseError(RuntimeError):
    def __init__(self, code: str, message: str):
        super().__init__(message)
        self.code = code


@dataclass
class ParsedElement:
    kind: str
    text: str
    section_path: list[str] = field(default_factory=list)
    location: dict[str, Any] = field(default_factory=dict)
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedDocument:
    parser: str
    elements: list[ParsedElement]
    page_count: int = 0
    sheet_count: int = 0
    metadata: dict[str, Any] = field(default_factory=dict)


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".txt", ".md", ".markdown", ".csv", ".tsv", ".json", ".jsonl"}


def _decode_bytes(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "big5"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentParseError("DOCUMENT_ENCODING_UNSUPPORTED", "无法识别文档编码")


def _clean(text: Any) -> str:
    return re.sub(r"[ \t\r\f\v]+", " ", str(text or "")).strip()


def _is_faq_header(value: str) -> bool:
    normalized = re.sub(r"\s+", "", value).lower()
    return normalized in {
        "标准问题",
        "问题",
        "question",
        "query",
        "标准答案",
        "答案",
        "answer",
    }


def _faq_row(row: dict[str, Any], index: int, source: str) -> ParsedElement:
    normalized = {str(k).strip().lower(): _clean(v) for k, v in row.items() if k is not None}

    def pick(*names: str) -> str:
        for name in names:
            for key, value in normalized.items():
                if key == name or key.replace("_", "") == name.replace("_", ""):
                    return value
        return ""

    question = pick("标准问题", "问题", "question", "query")
    answer = pick("标准答案", "答案", "answer", "response")
    if not question and not answer:
        text = "\n".join(f"{key}: {_clean(value)}" for key, value in row.items() if _clean(value))
        return ParsedElement("row", text, location={"row": index}, metadata={"source": source})
    fields = {
        "question": question,
        "answer": answer,
        "keywords": pick("关键词", "keywords", "key_words"),
        "category": pick("大类", "category"),
        "subcategory": pick("小类", "subcategory", "sub_category"),
    }
    text = "\n".join(
        f"{label}: {value}"
        for label, value in (
            ("问题", fields["question"]),
            ("答案", fields["answer"]),
            ("关键词", fields["keywords"]),
            ("大类", fields["category"]),
            ("小类", fields["subcategory"]),
        )
        if value
    )
    return ParsedElement(
        "faq_row",
        text,
        location={"row": index},
        metadata={"source": source, "structured": True, **fields},
    )


def parse_csv(path: Path, delimiter: str | None = None) -> ParsedDocument:
    text = _decode_bytes(path.read_bytes())
    if delimiter is None:
        try:
            delimiter = csv.Sniffer().sniff(text[:8192], delimiters=",\t;").delimiter
        except csv.Error:
            delimiter = "\t" if path.suffix.lower() == ".tsv" else ","
    reader = csv.DictReader(io.StringIO(text), delimiter=delimiter)
    if not reader.fieldnames:
        raise DocumentParseError("DOCUMENT_PARSE_FAILED", "CSV缺少表头")
    elements = [_faq_row(row, index + 2, path.name) for index, row in enumerate(reader)]
    elements = [element for element in elements if element.text]
    if not elements:
        raise DocumentParseError("DOCUMENT_EMPTY", "文档没有可索引内容")
    return ParsedDocument("csv_faq" if any(e.kind == "faq_row" for e in elements) else "csv", elements, metadata={"headers": reader.fieldnames})


def parse_json(path: Path) -> ParsedDocument:
    text = _decode_bytes(path.read_bytes())
    try:
        payload: Any
        if path.suffix.lower() == ".jsonl":
            payload = [json.loads(line) for line in text.splitlines() if line.strip()]
        else:
            payload = json.loads(text)
    except json.JSONDecodeError as exc:
        raise DocumentParseError("DOCUMENT_PARSE_FAILED", f"JSON解析失败: {exc.msg}") from exc
    rows = payload if isinstance(payload, list) else [payload]
    elements = []
    for index, row in enumerate(rows, start=1):
        if isinstance(row, dict):
            elements.append(_faq_row(row, index, path.name))
        else:
            elements.append(ParsedElement("paragraph", _clean(row), location={"item": index}))
    elements = [element for element in elements if element.text]
    if not elements:
        raise DocumentParseError("DOCUMENT_EMPTY", "文档没有可索引内容")
    return ParsedDocument("json", elements)


def parse_text(path: Path, markdown: bool = False) -> ParsedDocument:
    text = _decode_bytes(path.read_bytes())
    if not text.strip():
        raise DocumentParseError("DOCUMENT_EMPTY", "文档没有可索引内容")
    elements: list[ParsedElement] = []
    headings: list[str] = []
    buffer: list[str] = []
    start_line = 1

    def flush(end_line: int) -> None:
        nonlocal buffer, start_line
        value = _clean("\n".join(buffer))
        if value:
            elements.append(
                ParsedElement(
                    "paragraph",
                    value,
                    section_path=headings.copy(),
                    location={"lineStart": start_line, "lineEnd": end_line},
                )
            )
        buffer = []

    lines = text.splitlines()
    for line_no, line in enumerate(lines, start=1):
        heading_match = re.match(r"^\s{0,3}(#{1,6})\s+(.+?)\s*#*\s*$", line) if markdown else None
        if heading_match:
            flush(line_no - 1)
            level = len(heading_match.group(1))
            heading = _clean(heading_match.group(2))
            headings = headings[: level - 1] + [heading]
            elements.append(
                ParsedElement("heading", heading, section_path=headings.copy(), location={"line": line_no}, metadata={"level": level})
            )
        elif not line.strip():
            flush(line_no - 1)
            start_line = line_no + 1
        else:
            if not buffer:
                start_line = line_no
            buffer.append(line)
    flush(len(lines))
    if not elements:
        raise DocumentParseError("DOCUMENT_EMPTY", "文档没有可索引内容")
    return ParsedDocument("markdown" if markdown else "text", elements)


def parse_pdf(path: Path) -> ParsedDocument:
    try:
        import pymupdf as fitz
    except ImportError as exc:
        raise DocumentParseError("PARSER_DEPENDENCY_MISSING", "缺少PyMuPDF依赖") from exc
    try:
        document = fitz.open(path)
    except Exception as exc:  # pragma: no cover - library-specific errors
        raise DocumentParseError("DOCUMENT_PARSE_FAILED", "PDF无法打开") from exc
    elements: list[ParsedElement] = []
    text_char_count = 0
    for page_index, page in enumerate(document, start=1):
        blocks = page.get_text("blocks")
        page_text = "\n".join(block[4] for block in blocks if len(block) >= 5 and block[4].strip())
        text_char_count += len(page_text.strip())
        if page_text.strip():
            elements.append(
                ParsedElement(
                    "page",
                    _clean(page_text),
                    location={"page": page_index},
                )
            )
    page_count = len(document)
    document.close()
    if text_char_count == 0:
        raise DocumentParseError("SCANNED_DOCUMENT_UNSUPPORTED", "PDF没有可提取文本，V1不支持扫描件OCR")
    return ParsedDocument("pymupdf", elements, page_count=page_count)


def parse_docx(path: Path) -> ParsedDocument:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentParseError("PARSER_DEPENDENCY_MISSING", "缺少python-docx依赖") from exc
    try:
        document = Document(path)
    except Exception as exc:  # pragma: no cover
        raise DocumentParseError("DOCUMENT_PARSE_FAILED", "DOCX无法打开") from exc
    elements: list[ParsedElement] = []
    headings: list[str] = []
    ordinal = 0
    for paragraph in document.paragraphs:
        value = _clean(paragraph.text)
        if not value:
            continue
        ordinal += 1
        style = str(paragraph.style.name or "")
        if style.lower().startswith("heading"):
            match = re.search(r"(\d+)", style)
            level = int(match.group(1)) if match else 1
            headings = headings[: level - 1] + [value]
            elements.append(ParsedElement("heading", value, headings.copy(), {"paragraph": ordinal}, {"level": level}))
        else:
            elements.append(ParsedElement("paragraph", value, headings.copy(), {"paragraph": ordinal}))
    for table_index, table in enumerate(document.tables, start=1):
        rows = table.rows
        if not rows:
            continue
        headers = [_clean(cell.text) for cell in rows[0].cells]
        for row_index, row in enumerate(rows[1:], start=2):
            values = [_clean(cell.text) for cell in row.cells]
            mapping = {headers[i] if i < len(headers) and headers[i] else f"列{i + 1}": values[i] if i < len(values) else "" for i in range(len(values))}
            element = _faq_row(mapping, row_index, f"{path.name}#table-{table_index}")
            element.kind = "table_row" if element.kind != "faq_row" else element.kind
            element.section_path = headings.copy()
            element.location = {"table": table_index, "row": row_index}
            element.metadata["tableHeaders"] = headers
            elements.append(element)
    if not elements:
        raise DocumentParseError("DOCUMENT_EMPTY", "DOCX没有可索引内容")
    return ParsedDocument("python-docx", elements)


def parse_xlsx(path: Path) -> ParsedDocument:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise DocumentParseError("PARSER_DEPENDENCY_MISSING", "缺少openpyxl依赖") from exc
    try:
        workbook = load_workbook(path, read_only=True, data_only=True)
    except Exception as exc:  # pragma: no cover
        raise DocumentParseError("DOCUMENT_PARSE_FAILED", "XLSX无法打开") from exc
    elements: list[ParsedElement] = []
    sheet_count = 0
    for sheet in workbook.worksheets:
        sheet_count += 1
        rows = sheet.iter_rows(values_only=True)
        try:
            headers = [_clean(value) for value in next(rows)]
        except StopIteration:
            continue
        for row_index, row in enumerate(rows, start=2):
            values = [_clean(value) for value in row]
            mapping = {headers[i] if i < len(headers) and headers[i] else f"列{i + 1}": values[i] if i < len(values) else "" for i in range(len(values))}
            element = _faq_row(mapping, row_index, f"{path.name}#{sheet.title}")
            element.location = {"sheet": sheet.title, "row": row_index}
            element.metadata["sheet"] = sheet.title
            element.metadata["tableHeaders"] = headers
            elements.append(element)
    workbook.close()
    if not elements:
        raise DocumentParseError("DOCUMENT_EMPTY", "XLSX没有可索引内容")
    return ParsedDocument("openpyxl", elements, sheet_count=sheet_count)


def parse_document(path: Path) -> ParsedDocument:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise DocumentParseError("UNSUPPORTED_FILE_TYPE", f"不支持的文件类型: {suffix or 'unknown'}")
    if suffix == ".pdf":
        return parse_pdf(path)
    if suffix == ".docx":
        return parse_docx(path)
    if suffix == ".xlsx":
        return parse_xlsx(path)
    if suffix in {".csv", ".tsv"}:
        return parse_csv(path, "\t" if suffix == ".tsv" else None)
    if suffix in {".json", ".jsonl"}:
        return parse_json(path)
    return parse_text(path, markdown=suffix in {".md", ".markdown"})
