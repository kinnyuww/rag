from __future__ import annotations

from pathlib import Path

import pymupdf
import pytest
from docx import Document
from openpyxl import Workbook

from app.chunking import build_chunks
from app.parsers import DocumentParseError, parse_document


def test_text_markdown_docx_xlsx_and_pdf_parsers(tmp_path: Path):
    text_path = tmp_path / "source.txt"
    text_path.write_text("第一段内容。\n\n第二段内容。", encoding="utf-8")
    markdown_path = tmp_path / "source.md"
    markdown_path.write_text("# 标题\n\n说明内容。\n\n## 子标题\n\n更多内容。", encoding="utf-8")

    docx_path = tmp_path / "source.docx"
    docx = Document()
    docx.add_heading("办理说明", level=1)
    docx.add_paragraph("请携带有效证件。")
    table = docx.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "标准问题"
    table.rows[0].cells[1].text = "标准答案"
    table.rows[1].cells[0].text = "如何办理？"
    table.rows[1].cells[1].text = "到窗口办理。"
    docx.save(docx_path)

    xlsx_path = tmp_path / "source.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "FAQ"
    sheet.append(["标准问题", "标准答案", "关键词"])
    sheet.append(["在哪里办理？", "到政务窗口办理。", "政务,窗口"])
    workbook.save(xlsx_path)

    pdf_path = tmp_path / "source.pdf"
    pdf = pymupdf.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Text-bearing PDF source for local RAG.")
    pdf.save(pdf_path)
    pdf.close()

    parsed = [parse_document(path) for path in (text_path, markdown_path, docx_path, xlsx_path, pdf_path)]
    assert [item.parser for item in parsed] == ["text", "markdown", "python-docx", "openpyxl", "pymupdf"]
    assert all(item.elements for item in parsed)
    for index, item in enumerate(parsed):
        _, chunks = build_chunks(item, document_id=f"doc-{index}", document_version=1)
        assert chunks
        assert all(chunk["content_hash"] for chunk in chunks)


def test_blank_pdf_requires_ocr_and_fails_explicitly(tmp_path: Path):
    path = tmp_path / "scan.pdf"
    document = pymupdf.open()
    document.new_page()
    document.save(path)
    document.close()
    with pytest.raises(DocumentParseError) as caught:
        parse_document(path)
    assert caught.value.code == "SCANNED_DOCUMENT_UNSUPPORTED"
